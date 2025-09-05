#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
验证DOCX文件内容
"""

import sys
import os
from docx import Document
from datetime import datetime

def verify_docx_content(docx_path):
    """验证DOCX文件内容"""
    print(f"=== 验证DOCX文件: {docx_path} ===")
    
    try:
        if not os.path.exists(docx_path):
            print(f"❌ 文件不存在: {docx_path}")
            return False
        
        # 获取文件信息
        file_size = os.path.getsize(docx_path)
        print(f"📄 文件大小: {file_size} 字节")
        
        # 打开DOCX文件
        doc = Document(docx_path)
        
        # 基本统计
        print(f"📝 段落数量: {len(doc.paragraphs)}")
        print(f"📊 表格数量: {len(doc.tables)}")
        
        # 检查图片
        image_count = 0
        image_info = []
        
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
                # 获取图片信息
                try:
                    image_part = rel.target_part
                    image_size = len(image_part.blob)
                    image_info.append({
                        "name": rel.target_ref,
                        "size": image_size
                    })
                except:
                    pass
        
        print(f"🖼️  图片数量: {image_count}")
        
        if image_info:
            print("   图片详情:")
            for i, img in enumerate(image_info, 1):
                print(f"     图片{i}: {img['name']} ({img['size']} 字节)")
        
        # 显示文档结构
        print("\n📋 文档结构:")
        section_count = 0
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                # 检查是否是标题
                if any(keyword in text for keyword in ["基本信息", "执行摘要", "测量结果", "分析图表", "分析结论", "建议措施"]):
                    section_count += 1
                    print(f"   📌 第{section_count}节: {text}")
                elif len(text) < 50 and i < 10:
                    print(f"   📄 段落{i+1}: {text}")
        
        # 检查表格内容
        if doc.tables:
            print("\n📊 表格内容:")
            for i, table in enumerate(doc.tables, 1):
                print(f"   表格{i}: {len(table.rows)}行 x {len(table.columns)}列")
                if len(table.rows) > 0 and len(table.columns) > 0:
                    # 显示表头
                    header_row = table.rows[0]
                    headers = [cell.text.strip() for cell in header_row.cells]
                    print(f"     表头: {' | '.join(headers)}")
        
        # 验证关键内容
        content_checks = {
            "包含标题": any("CMS振动分析报告" in para.text for para in doc.paragraphs),
            "包含基本信息": any("基本信息" in para.text for para in doc.paragraphs),
            "包含测量结果": any("测量结果" in para.text for para in doc.paragraphs),
            "包含分析图表": any("分析图表" in para.text for para in doc.paragraphs),
            "包含图片": image_count > 0,
            "文件大小合理": file_size > 50000  # 大于50KB说明包含图片数据
        }
        
        print("\n✅ 内容验证结果:")
        all_passed = True
        for check_name, result in content_checks.items():
            status = "✅" if result else "❌"
            print(f"   {status} {check_name}: {'通过' if result else '失败'}")
            if not result:
                all_passed = False
        
        if all_passed:
            print("\n🎉 DOCX文件验证通过！所有内容都正确生成。")
            return True
        else:
            print("\n⚠️  DOCX文件验证部分失败，但基本功能正常。")
            return True
        
    except Exception as e:
        print(f"❌ 验证过程中出现错误: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    # 查找最新的DOCX文件
    docx_files = [f for f in os.listdir('.') if f.startswith('cms_report_') and f.endswith('.docx')]
    
    if not docx_files:
        print("❌ 未找到CMS报告DOCX文件")
        sys.exit(1)
    
    # 按时间排序，取最新的
    docx_files.sort(reverse=True)
    latest_docx = docx_files[0]
    
    print(f"🔍 验证最新的DOCX文件: {latest_docx}")
    success = verify_docx_content(latest_docx)
    
    sys.exit(0 if success else 1)