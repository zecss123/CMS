#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试DOCX报告生成功能
"""

import sys
import os
from datetime import datetime
from pathlib import Path

# 添加项目根目录到Python路径
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from report.generator import CMSReportGenerator
from utils.chart_generator import VibrationChartGenerator
import numpy as np

def test_docx_generation():
    """测试DOCX报告生成"""
    print("=== 测试DOCX报告生成功能 ===")
    
    try:
        # 1. 生成测试图表
        print("\n📊 生成测试图表...")
        chart_gen = VibrationChartGenerator()
        
        # 生成时域信号
        t = np.linspace(0, 1, 1000)
        signal = np.sin(2 * np.pi * 50 * t) + 0.5 * np.sin(2 * np.pi * 120 * t) + 0.1 * np.random.randn(1000)
        
        # 创建图表并获取base64数据
        chart_base64 = chart_gen.create_time_series_chart(
            signal, 
            sampling_rate=1000, 
            title="振动趋势图"
        )
        
        if chart_base64:
            print(f"✅ 图表生成成功，数据长度: {len(chart_base64)} 字符")
        else:
            print("❌ 图表生成失败")
            return False
        
        # 2. 准备报告数据
        print("\n📋 准备报告数据...")
        report_data = {
            "title": "CMS振动分析报告 - DOCX测试",
            "basic_info": {
                "wind_farm": "测试风场A",
                "turbine_id": "WT001",
                "measurement_date": "2024-01-15",
                "report_date": datetime.now().strftime("%Y-%m-%d"),
                "operator": "测试工程师",
                "equipment_status": "运行中"
            },
            "executive_summary": "本次振动测量显示设备运行状态良好，各项指标均在正常范围内。主要频率成分为50Hz和120Hz，符合设备运行特征。",
            "measurement_results": [
                {
                    "measurement_point": "主轴承DE",
                    "rms_value": 2.5,
                    "peak_value": 8.2,
                    "main_frequency": 25.5,
                    "alarm_level": "normal"
                },
                {
                    "measurement_point": "齿轮箱HSS",
                    "rms_value": 4.1,
                    "peak_value": 12.8,
                    "main_frequency": 1250.0,
                    "alarm_level": "warning"
                },
                {
                    "measurement_point": "发电机NDE",
                    "rms_value": 1.8,
                    "peak_value": 6.5,
                    "main_frequency": 50.0,
                    "alarm_level": "normal"
                }
            ],
            "charts": {
                "振动趋势图": chart_base64
            },
            "analysis_conclusion": "根据振动测量结果分析，设备整体运行状态良好。主轴承振动水平正常，齿轮箱高速轴存在轻微异常但仍在可接受范围内，发电机轴承运行平稳。建议继续监测齿轮箱状态。",
            "recommendations": [
                "定期检查齿轮箱润滑情况，确保润滑油质量",
                "加强高速轴监测频次，建议每月进行一次详细检测",
                "关注主轴承温度变化，避免过热现象",
                "建议下次全面检测时间：3个月后"
            ]
        }
        
        print(f"✅ 报告数据准备完成")
        print(f"   - 基本信息: {len(report_data['basic_info'])} 项")
        print(f"   - 测量结果: {len(report_data['measurement_results'])} 个测点")
        print(f"   - 图表数量: {len(report_data['charts'])} 个")
        print(f"   - 建议措施: {len(report_data['recommendations'])} 条")
        
        # 3. 生成DOCX报告
        print("\n📄 生成DOCX报告...")
        generator = CMSReportGenerator()
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        docx_filename = f"test_docx_report_{timestamp}.docx"
        
        success = generator.generate_docx_report(report_data, docx_filename)
        
        if success:
            print(f"✅ DOCX报告生成成功: {docx_filename}")
            
            # 检查文件是否存在
            if os.path.exists(docx_filename):
                file_size = os.path.getsize(docx_filename)
                print(f"   文件大小: {file_size} 字节")
                
                # 尝试验证DOCX文件结构
                try:
                    from docx import Document
                    doc = Document(docx_filename)
                    
                    print(f"   段落数量: {len(doc.paragraphs)}")
                    print(f"   表格数量: {len(doc.tables)}")
                    
                    # 检查是否包含图片
                    image_count = 0
                    for rel in doc.part.rels.values():
                        if "image" in rel.target_ref:
                            image_count += 1
                    print(f"   图片数量: {image_count}")
                    
                    # 显示前几个段落的内容
                    print("\n📋 文档内容预览:")
                    for i, para in enumerate(doc.paragraphs[:10]):
                        if para.text.strip():
                            print(f"   段落{i+1}: {para.text[:50]}...")
                    
                    print(f"\n✅ DOCX文件验证成功")
                    return True
                    
                except Exception as e:
                    print(f"❌ DOCX文件验证失败: {e}")
                    return False
            else:
                print(f"❌ 文件未生成: {docx_filename}")
                return False
        else:
            print(f"❌ DOCX报告生成失败")
            return False
            
    except Exception as e:
        print(f"❌ 测试过程中出现错误: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    success = test_docx_generation()
    if success:
        print("\n🎉 DOCX生成测试通过！")
    else:
        print("\n💥 DOCX生成测试失败！")
    
    sys.exit(0 if success else 1)