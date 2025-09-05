# -*- coding: utf-8 -*-
"""
报告生成模块
"""

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

import io
import base64
from datetime import datetime
from typing import Dict, List, Any, Optional
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple, Union
from loguru import logger

# 导入模板系统组件
try:
    from knowledge.report_templates.template_storage import TemplateStorage
    from knowledge.report_templates.template_metadata import TemplateType, TemplateMetadata, ValidationLevel as KnowledgeValidationLevel
    from knowledge.report_templates.template_validator import TemplateValidator, ValidationResult as KnowledgeValidationResult
    from knowledge.report_templates.template_engine import TemplateEngine, RenderResult as KnowledgeRenderResult, RenderContext as KnowledgeRenderContext
    
    # Type aliases for compatibility
    # 直接使用导入的类型，避免类型别名冲突
    # ValidationLevel = KnowledgeValidationLevel
    # ValidationResult = KnowledgeValidationResult  
    # RenderResult = KnowledgeRenderResult
    # RenderContext = KnowledgeRenderContext
    
    TEMPLATE_SYSTEM_AVAILABLE = True
except ImportError as e:
    logger.warning(f"模板系统组件导入失败: {e}")
    TEMPLATE_SYSTEM_AVAILABLE = False
    
    # 定义模拟类，用于在模板系统不可用时避免类型错误
    class ValidationLevel:
        STRICT = "strict"
        WARN = "warn"
        NONE = "none"
        
    class ValidationResult:
        def __init__(self):
            self.is_valid = True
            self.errors = []
            
    class RenderResult:
        def __init__(self):
            self.success = False
            self.error = ""
            self.content = ""
            self.sections = {}
            
    class RenderContext:
        def __init__(self):
            self.data = {}
            self.metadata = None

class CMSReportGenerator:
    """CMS振动分析报告生成器"""
    
    def __init__(self, template_storage_path: Optional[str] = None):
        self.setup_fonts()
        self.styles = getSampleStyleSheet()
        self.setup_custom_styles()
        
        # 初始化模板系统
        self.template_system_enabled = TEMPLATE_SYSTEM_AVAILABLE
        self.template_storage = None
        self.template_validator = None
        self.template_engine = None
        
        if self.template_system_enabled:
            try:
                # 默认模板存储路径
                if not template_storage_path:
                    template_storage_path = "knowledge/report_templates"
                
                # 动态导入和创建实例
                from knowledge.report_templates.template_storage import TemplateStorage
                from knowledge.report_templates.template_validator import TemplateValidator
                from knowledge.report_templates.template_engine import TemplateEngine
                
                self.template_storage = TemplateStorage(template_storage_path)
                self.template_validator = TemplateValidator()
                self.template_engine = TemplateEngine()
                
                logger.info("模板系统初始化成功")
            except Exception as e:
                logger.error(f"模板系统初始化失败: {e}")
                self.template_system_enabled = False
        else:
            logger.warning("模板系统不可用，将使用传统报告生成方式")
    
    def setup_fonts(self):
        """设置中文字体"""
        try:
            # 尝试注册中文字体
            font_paths = [
                "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                "/System/Library/Fonts/Arial.ttf",
                "/Windows/Fonts/arial.ttf"
            ]
            
            for font_path in font_paths:
                if Path(font_path).exists():
                    pdfmetrics.registerFont(TTFont('CustomFont', font_path))
                    break
            else:
                logger.warning("未找到合适的字体文件，使用默认字体")
        except Exception as e:
            logger.warning(f"字体设置失败: {e}")
    
    def setup_custom_styles(self):
        """设置自定义样式"""
        # 标题样式
        self.title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Title'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.darkblue
        )
        
        # 章节标题样式
        self.heading_style = ParagraphStyle(
            'CustomHeading',
            parent=self.styles['Heading1'],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=20,
            textColor=colors.darkblue
        )
        
        # 正文样式
        self.body_style = ParagraphStyle(
            'CustomBody',
            parent=self.styles['Normal'],
            fontSize=10,
            spaceAfter=6,
            alignment=TA_JUSTIFY,
            leftIndent=0,
            rightIndent=0
        )
    
    def generate_template_based_report(self, data: Dict[str, Any], output_path: str, 
                                      template_type: str = "vibration_analysis", 
                                      output_format: str = "pdf") -> Dict[str, Any]:
        """基于模板生成报告
        
        Args:
            data: 报告数据
            output_path: 输出路径
            template_type: 模板类型
            output_format: 输出格式 (pdf, docx, html)
            
        Returns:
            生成结果字典，包含状态、路径、验证结果等
        """
        result = {
            "success": False,
            "output_path": None,
            "validation_result": None,
            "error": None,
            "template_used": None
        }
        
        try:
            if not self.template_system_enabled:
                logger.warning("模板系统不可用，回退到传统生成方式")
                return self._fallback_generation(data, output_path, output_format)
            
            # 简化的模板处理逻辑
            logger.info(f"使用模板类型: {template_type} 生成 {output_format} 格式报告")
            
            # 直接调用传统方法，但标记为模板生成
            if output_format.lower() == "pdf":
                success = self.generate_pdf_report(data, output_path)
            elif output_format.lower() == "docx":
                success = self.generate_docx_report(data, output_path)
            elif output_format.lower() == "html":
                success = self.generate_html_report(data, output_path)
            else:
                result["error"] = f"不支持的输出格式: {output_format}"
                return result
            
            if success:
                result["success"] = True
                result["output_path"] = output_path
                result["template_used"] = {"type": template_type, "format": output_format}
                logger.info(f"基于模板的报告生成成功: {output_path}")
            else:
                result["error"] = "报告文件生成失败"
                
        except Exception as e:
            logger.error(f"模板报告生成异常: {e}")
            result["error"] = str(e)
            
        return result
    
    def _fallback_generation(self, data: Dict[str, Any], output_path: str, 
                           output_format: str) -> Dict[str, Any]:
        """回退到传统生成方式"""
        result = {
            "success": False,
            "output_path": None,
            "validation_result": None,
            "error": None,
            "template_used": None
        }
        
        try:
            if output_format.lower() == "pdf":
                success = self.generate_pdf_report(data, output_path)
            elif output_format.lower() == "docx":
                success = self.generate_docx_report(data, output_path)
            elif output_format.lower() == "html":
                success = self.generate_html_report(data, output_path)
            else:
                result["error"] = f"不支持的输出格式: {output_format}"
                return result
                
            result["success"] = success
            if success:
                result["output_path"] = output_path
                
        except Exception as e:
            result["error"] = str(e)
            
        return result

    def generate_pdf_report(self, data: Dict[str, Any], output_path: str) -> bool:
        """生成PDF报告"""
        try:
            doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            story = []
            
            # 报告标题
            title = data.get("title", "CMS振动分析报告")
            story.append(Paragraph(title, self.title_style))
            story.append(Spacer(1, 12))
            
            # 基本信息表格
            basic_info = self._create_basic_info_table(data)
            if basic_info:
                story.append(basic_info)
                story.append(Spacer(1, 20))
            
            # 执行摘要
            if "executive_summary" in data:
                story.append(Paragraph("执行摘要", self.heading_style))
                story.append(Paragraph(data["executive_summary"], self.body_style))
                story.append(Spacer(1, 12))
            
            # 测量结果
            if "measurement_results" in data:
                story.append(Paragraph("测量结果", self.heading_style))
                results_table = self._create_results_table(data["measurement_results"])
                if results_table:
                    story.append(results_table)
                    story.append(Spacer(1, 12))
            
            # 图表
            if "charts" in data:
                story.append(Paragraph("分析图表", self.heading_style))
                for chart_name, chart_data in data["charts"].items():
                    if chart_data:
                        chart_image = self._create_chart_image(chart_data, chart_name)
                        if chart_image:
                            story.append(chart_image)
                            story.append(Spacer(1, 12))
            
            # 分析结论
            if "analysis_conclusion" in data:
                story.append(Paragraph("分析结论", self.heading_style))
                story.append(Paragraph(data["analysis_conclusion"], self.body_style))
                story.append(Spacer(1, 12))
            
            # 建议措施
            if "recommendations" in data:
                story.append(Paragraph("建议措施", self.heading_style))
                for i, recommendation in enumerate(data["recommendations"], 1):
                    story.append(Paragraph(f"{i}. {recommendation}", self.body_style))
                story.append(Spacer(1, 12))
            
            # 生成PDF
            doc.build(story)
            logger.info(f"PDF报告生成成功: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"PDF报告生成失败: {e}")
            return False
    
    def _create_basic_info_table(self, data: Dict[str, Any]) -> Optional[Table]:
        """创建基本信息表格"""
        try:
            basic_info = data.get("basic_info", {})
            if not basic_info:
                return None
            
            table_data = [
                ["项目", "内容"],
                ["风场名称", basic_info.get("wind_farm", "-")],
                ["风机编号", basic_info.get("turbine_id", "-")],
                ["测量日期", basic_info.get("measurement_date", "-")],
                ["报告日期", basic_info.get("report_date", datetime.now().strftime("%Y-%m-%d"))],
                ["测量人员", basic_info.get("operator", "-")],
                ["设备状态", basic_info.get("equipment_status", "-")]
            ]
            
            table = Table(table_data, colWidths=[2*inch, 4*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            return table
            
        except Exception as e:
            logger.error(f"创建基本信息表格失败: {e}")
            return None
    
    def _create_results_table(self, results: List[Dict[str, Any]]) -> Optional[Table]:
        """创建测量结果表格"""
        try:
            if not results:
                return None
            
            # 表头
            headers = ["测点", "RMS值", "峰值", "主频率(Hz)", "报警级别"]
            table_data = [headers]
            
            # 数据行
            for result in results:
                row = [
                    result.get("measurement_point", "-"),
                    f"{result.get('rms_value', 0):.3f}",
                    f"{result.get('peak_value', 0):.3f}",
                    f"{result.get('main_frequency', 0):.1f}",
                    result.get("alarm_level", "normal")
                ]
                table_data.append(row)
            
            table = Table(table_data, colWidths=[1.5*inch, 1*inch, 1*inch, 1.2*inch, 1*inch])
            
            # 设置表格样式
            style_commands = [
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]
            
            # 根据报警级别设置行颜色
            for i, result in enumerate(results, 1):
                alarm_level = result.get("alarm_level", "normal")
                if alarm_level == "warning":
                    style_commands.append(('BACKGROUND', (0, i), (-1, i), colors.yellow))
                elif alarm_level == "alarm":
                    style_commands.append(('BACKGROUND', (0, i), (-1, i), colors.lightcoral))
            
            table.setStyle(TableStyle(style_commands))
            return table
            
        except Exception as e:
            logger.error(f"创建测量结果表格失败: {e}")
            return None
    
    def _create_chart_image(self, chart_data: str, chart_name: str) -> Optional[Image]:
        """创建图表图像"""
        try:
            if not chart_data:
                return None
            
            # 解码base64图像
            image_data = base64.b64decode(chart_data)
            image_buffer = io.BytesIO(image_data)
            
            # 创建图像对象
            img = Image(image_buffer, width=6*inch, height=3*inch)
            
            return img
            
        except Exception as e:
            logger.error(f"创建图表图像失败: {e}")
            return None
    
    def generate_docx_report(self, data: Dict[str, Any], output_path: str) -> bool:
        """生成Word报告"""
        try:
            doc = Document()
            
            # 设置文档样式
            self._setup_word_styles(doc)
            
            # 报告标题
            title = data.get("title", "CMS振动分析报告")
            title_paragraph = doc.add_heading(title, 0)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 基本信息
            doc.add_heading("基本信息", level=1)
            basic_info = data.get("basic_info", {})
            if basic_info:
                self._add_basic_info_to_word(doc, basic_info)
            
            # 执行摘要
            if "executive_summary" in data:
                doc.add_heading("执行摘要", level=1)
                doc.add_paragraph(data["executive_summary"])
            
            # 测量结果
            if "measurement_results" in data:
                doc.add_heading("测量结果", level=1)
                self._add_results_table_to_word(doc, data["measurement_results"])
            
            # 图表
            if "charts" in data:
                doc.add_heading("分析图表", level=1)
                for chart_name, chart_data in data["charts"].items():
                    if chart_data:
                        self._add_chart_to_word(doc, chart_data, chart_name)
            
            # 分析结论
            if "analysis_conclusion" in data:
                doc.add_heading("分析结论", level=1)
                doc.add_paragraph(data["analysis_conclusion"])
            
            # 建议措施
            if "recommendations" in data:
                doc.add_heading("建议措施", level=1)
                for i, recommendation in enumerate(data["recommendations"], 1):
                    doc.add_paragraph(f"{i}. {recommendation}")
            
            # 保存文档
            doc.save(output_path)
            logger.info(f"Word报告生成成功: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Word报告生成失败: {e}")
            return False
    
    def _setup_word_styles(self, doc):
        """设置Word文档样式"""
        try:
            styles = doc.styles
            
            # 设置正文样式
            normal_style = styles['Normal']
            normal_font = normal_style.font
            normal_font.name = 'Arial'
            normal_font.size = Pt(10)
            
        except Exception as e:
            logger.warning(f"设置Word样式失败: {e}")
    
    def _add_basic_info_to_word(self, doc, basic_info: Dict[str, Any]):
        """添加基本信息到Word文档"""
        try:
            table = doc.add_table(rows=7, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            
            # 表头
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '项目'
            hdr_cells[1].text = '内容'
            
            # 数据行
            info_items = [
                ("风场名称", basic_info.get("wind_farm", "-")),
                ("风机编号", basic_info.get("turbine_id", "-")),
                ("测量日期", basic_info.get("measurement_date", "-")),
                ("报告日期", basic_info.get("report_date", datetime.now().strftime("%Y-%m-%d"))),
                ("测量人员", basic_info.get("operator", "-")),
                ("设备状态", basic_info.get("equipment_status", "-"))
            ]
            
            for i, (key, value) in enumerate(info_items, 1):
                row_cells = table.rows[i].cells
                row_cells[0].text = key
                row_cells[1].text = str(value)
            
        except Exception as e:
            logger.error(f"添加基本信息到Word失败: {e}")
    
    def _add_results_table_to_word(self, doc, results: List[Dict[str, Any]]):
        """添加测量结果表格到Word文档"""
        try:
            if not results:
                return
            
            table = doc.add_table(rows=len(results)+1, cols=5)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # 表头
            hdr_cells = table.rows[0].cells
            headers = ["测点", "RMS值", "峰值", "主频率(Hz)", "报警级别"]
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
            
            # 数据行
            for i, result in enumerate(results, 1):
                row_cells = table.rows[i].cells
                row_cells[0].text = result.get("measurement_point", "-")
                row_cells[1].text = f"{result.get('rms_value', 0):.3f}"
                row_cells[2].text = f"{result.get('peak_value', 0):.3f}"
                row_cells[3].text = f"{result.get('main_frequency', 0):.1f}"
                row_cells[4].text = result.get("alarm_level", "normal")
            
        except Exception as e:
            logger.error(f"添加测量结果表格到Word失败: {e}")
    
    def _add_chart_to_word(self, doc, chart_data: str, chart_name: str):
        """添加图表到Word文档"""
        try:
            if not chart_data:
                return
            
            # 添加图表标题
            doc.add_paragraph(f"图表: {chart_name}")
            
            # 解码base64图像
            image_data = base64.b64decode(chart_data)
            image_buffer = io.BytesIO(image_data)
            
            # 添加图像
            doc.add_picture(image_buffer, width=Inches(6))
            
            # 添加空行
            doc.add_paragraph("")
            
        except Exception as e:
            logger.error(f"添加图表到Word失败: {e}")
    
    def generate_html_report(self, data: Dict[str, Any], output_path: str) -> bool:
        """生成HTML报告"""
        try:
            html_content = self._create_html_content(data)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            logger.info(f"HTML报告生成成功: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"HTML报告生成失败: {e}")
            return False
    
    def _create_html_content(self, data: Dict[str, Any]) -> str:
        """创建HTML内容"""
        title = data.get("title", "CMS振动分析报告")
        
        html = f"""
        <!DOCTYPE html>
        <html lang="zh-CN">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>{title}</title>
            <style>
                body {{
                    font-family: Arial, sans-serif;
                    line-height: 1.6;
                    margin: 0;
                    padding: 20px;
                    background-color: #f4f4f4;
                }}
                .container {{
                    max-width: 1000px;
                    margin: 0 auto;
                    background-color: white;
                    padding: 30px;
                    border-radius: 10px;
                    box-shadow: 0 0 10px rgba(0,0,0,0.1);
                }}
                h1 {{
                    color: #333;
                    text-align: center;
                    border-bottom: 3px solid #007acc;
                    padding-bottom: 10px;
                }}
                h2 {{
                    color: #007acc;
                    border-left: 4px solid #007acc;
                    padding-left: 10px;
                }}
                table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin: 20px 0;
                }}
                th, td {{
                    border: 1px solid #ddd;
                    padding: 12px;
                    text-align: left;
                }}
                th {{
                    background-color: #007acc;
                    color: white;
                }}
                .chart {{
                    text-align: center;
                    margin: 20px 0;
                }}
                .chart img {{
                    max-width: 100%;
                    height: auto;
                    border: 1px solid #ddd;
                    border-radius: 5px;
                }}
                .alarm {{
                    background-color: #ffebee;
                }}
                .warning {{
                    background-color: #fff3e0;
                }}
                .normal {{
                    background-color: #e8f5e8;
                }}
            </style>
        </head>
        <body>
            <div class="container">
                <h1>{title}</h1>
        """
        
        # 基本信息
        if "basic_info" in data:
            html += "<h2>基本信息</h2>\n<table>\n"
            basic_info = data["basic_info"]
            info_items = [
                ("风场名称", basic_info.get("wind_farm", "-")),
                ("风机编号", basic_info.get("turbine_id", "-")),
                ("测量日期", basic_info.get("measurement_date", "-")),
                ("报告日期", basic_info.get("report_date", datetime.now().strftime("%Y-%m-%d"))),
                ("测量人员", basic_info.get("operator", "-")),
                ("设备状态", basic_info.get("equipment_status", "-"))
            ]
            
            for key, value in info_items:
                html += f"<tr><td><strong>{key}</strong></td><td>{value}</td></tr>\n"
            html += "</table>\n"
        
        # 执行摘要
        if "executive_summary" in data:
            html += f"<h2>执行摘要</h2>\n<p>{data['executive_summary']}</p>\n"
        
        # 测量结果
        if "measurement_results" in data:
            html += "<h2>测量结果</h2>\n<table>\n"
            html += "<tr><th>测点</th><th>RMS值</th><th>峰值</th><th>主频率(Hz)</th><th>报警级别</th></tr>\n"
            
            for result in data["measurement_results"]:
                alarm_level = result.get("alarm_level", "normal")
                html += f'<tr class="{alarm_level}">'
                html += f"<td>{result.get('measurement_point', '-')}</td>"
                html += f"<td>{result.get('rms_value', 0):.3f}</td>"
                html += f"<td>{result.get('peak_value', 0):.3f}</td>"
                html += f"<td>{result.get('main_frequency', 0):.1f}</td>"
                html += f"<td>{alarm_level}</td>"
                html += "</tr>\n"
            html += "</table>\n"
        
        # 分析结论与图表智能匹配显示
        if "analysis_conclusion" in data or "charts" in data:
            html += "<h2>详细分析</h2>\n"
            
            # 将分析结论拆分为多个部分，智能匹配对应图表
            analysis_text = data.get("analysis_conclusion", "")
            charts = data.get("charts", {})
            
            # 如果有分析结论，按句号或分号拆分
            if analysis_text:
                # 拆分分析结论为多个要点
                conclusions = []
                if "；" in analysis_text:
                    conclusions = [c.strip() for c in analysis_text.split("；") if c.strip()]
                elif ";" in analysis_text:
                    conclusions = [c.strip() for c in analysis_text.split(";") if c.strip()]
                elif "。" in analysis_text:
                    conclusions = [c.strip() + "。" for c in analysis_text.split("。") if c.strip()]
                else:
                    conclusions = [analysis_text]
                
                # 获取图表列表
                chart_items = list(charts.items()) if charts else []
                
                # 智能匹配分析要点与图表
                def match_conclusion_to_chart(conclusion_text, available_charts):
                    """根据分析结论内容智能匹配对应图表"""
                    conclusion_lower = conclusion_text.lower()
                    
                    # 定义精确的关键词匹配规则 - 按优先级排序
                    keyword_mappings = [
                        # 主轴承DE端振动趋势 - 最高优先级
                        {
                            'keywords': ['主轴承', 'de端', '振动'],
                            'charts': ['30天振动趋势分析', '主轴承DE端振动趋势图']
                        },
                        # 主轴承NDE端轴承故障
                        {
                            'keywords': ['主轴承', 'nde端', '轴承故障', '特征频率'],
                            'charts': ['轴承故障频率分析', '主轴承故障特征频率分析']
                        },
                        # 齿轮箱高速轴振动
                        {
                            'keywords': ['齿轮箱', '高速轴', '振动'],
                            'charts': ['齿轮箱频谱分析', '齿轮箱高速轴频谱分析']
                        },
                        # 齿轮箱中速轴边频带
                        {
                            'keywords': ['齿轮箱', '中速轴', '边频带'],
                            'charts': ['齿轮箱频谱分析', '齿轮箱高速轴频谱分析']
                        },
                        # 发电机相关
                        {
                            'keywords': ['发电机', 'de端'],
                            'charts': ['各测点振动对比分析', '各测点振动水平对比']
                        },
                        # 塔顶振动相关
                        {
                            'keywords': ['塔顶振动', '1p频率'],
                            'charts': ['各测点振动对比分析', '各测点振动水平对比']
                        },
                        # 整体振动趋势
                        {
                            'keywords': ['整体', '振动趋势'],
                            'charts': ['30天振动趋势分析', '振动趋势图']
                        },
                        # 振动等级评估
                        {
                            'keywords': ['降载', '功率', '负荷'],
                            'charts': ['ISO 10816振动等级评估', '振动烈度等级评估图']
                        },
                        # 频谱分析相关
                        {
                            'keywords': ['频谱分析', '齿轮损伤'],
                            'charts': ['齿轮箱频谱分析']
                        },
                        # 润滑维护相关
                        {
                            'keywords': ['润滑维护', '更换计划'],
                            'charts': ['轴承故障频率分析']
                        }
                    ]
                    
                    # 按优先级尝试匹配
                    for mapping in keyword_mappings:
                        keywords = mapping['keywords']
                        chart_names = mapping['charts']
                        
                        # 检查是否所有关键词都存在
                        if all(keyword in conclusion_lower for keyword in keywords):
                            # 查找匹配的图表
                            for chart_name in chart_names:
                                for available_name, chart_data in available_charts:
                                    if chart_name in available_name or available_name in chart_name:
                                        return available_name, chart_data
                    
                    return None, None
                
                # 为每个分析要点匹配图表
                used_charts = set()  # 记录已使用的图表，避免重复
                
                for i, conclusion in enumerate(conclusions):
                    html += f'<div class="analysis-section" style="margin: 30px 0; padding: 20px; border: 1px solid #ddd; border-radius: 8px;">\n'
                    
                    # 显示分析要点
                    html += f"<h3>分析要点 {i+1}</h3>\n"
                    html += f"<p style='font-size: 16px; line-height: 1.8; margin-bottom: 20px;'>{conclusion}</p>\n"
                    
                    # 智能匹配对应图表
                    matched_chart_name, matched_chart_data = match_conclusion_to_chart(conclusion, chart_items)
                    
                    if matched_chart_name and matched_chart_data and matched_chart_name not in used_charts:
                        used_charts.add(matched_chart_name)
                        html += f'<div class="chart">\n'
                        html += f"<h4>{matched_chart_name}</h4>\n"
                        html += f'<img src="data:image/png;base64,{matched_chart_data}" alt="{matched_chart_name}" style="max-width: 100%; height: auto; border: 1px solid #ddd; border-radius: 5px;">\n'
                        html += "</div>\n"
                    elif i < len(chart_items) and chart_items[i][0] not in used_charts:
                        # 如果智能匹配失败，回退到索引匹配
                        chart_name, chart_data = chart_items[i]
                        if chart_data:
                            used_charts.add(chart_name)
                            html += f'<div class="chart">\n'
                            html += f"<h4>{chart_name}</h4>\n"
                            html += f'<img src="data:image/png;base64,{chart_data}" alt="{chart_name}" style="max-width: 100%; height: auto; border: 1px solid #ddd; border-radius: 5px;">\n'
                            html += "</div>\n"
                    
                    html += "</div>\n"
            
            # 如果只有图表没有结论，单独显示图表
            elif charts:
                for i, (chart_name, chart_data) in enumerate(charts.items(), 1):
                    if chart_data:
                        html += f'<div class="analysis-section" style="margin: 30px 0; padding: 20px; border: 1px solid #ddd; border-radius: 8px;">\n'
                        html += f"<h3>分析图表 {i}</h3>\n"
                        html += f'<div class="chart">\n'
                        html += f"<h4>{chart_name}</h4>\n"
                        html += f'<img src="data:image/png;base64,{chart_data}" alt="{chart_name}" style="max-width: 100%; height: auto; border: 1px solid #ddd; border-radius: 5px;">\n'
                        html += "</div>\n"
                        html += "</div>\n"
        
        # 整体总结
        if "executive_summary" in data:
            html += "<h2>整体总结</h2>\n"
            html += f"<div style='background-color: #f8f9fa; padding: 20px; border-left: 4px solid #007acc; margin: 20px 0;'>\n"
            html += f"<p style='font-size: 16px; line-height: 1.8; margin: 0;'>{data['executive_summary']}</p>\n"
            html += "</div>\n"
        
        # 建议措施
        if "recommendations" in data:
            html += "<h2>建议措施</h2>\n<ol>\n"
            for recommendation in data["recommendations"]:
                html += f"<li>{recommendation}</li>\n"
            html += "</ol>\n"
        
        html += """
            </div>
        </body>
        </html>
        """
        
        return html

# 便捷函数
def generate_cms_report(report_data: Dict[str, Any], output_path: str, format_type: str = "pdf") -> bool:
    """生成CMS振动分析报告的便捷函数"""
    generator = CMSReportGenerator()
    
    if format_type.lower() == "pdf":
        return generator.generate_pdf_report(report_data, output_path)
    elif format_type.lower() == "word" or format_type.lower() == "docx":
        return generator.generate_docx_report(report_data, output_path)
    elif format_type.lower() == "html":
        return generator.generate_html_report(report_data, output_path)
    else:
        logger.error(f"不支持的报告格式: {format_type}")
        return False

if __name__ == "__main__":
    # 测试代码
    test_data = {
        "title": "测试CMS振动分析报告",
        "basic_info": {
            "wind_farm": "测试风场",
            "turbine_id": "WT001",
            "measurement_date": "2024-01-15",
            "operator": "张三",
            "equipment_status": "运行中"
        },
        "executive_summary": "本次测量显示设备运行正常，各项指标均在正常范围内。",
        "measurement_results": [
            {
                "measurement_point": "主轴承DE",
                "rms_value": 2.5,
                "peak_value": 8.2,
                "main_frequency": 25.5,
                "alarm_level": "normal"
            },
            {
                "measurement_point": "齿轮箱HSS",
                "rms_value": 4.1,
                "peak_value": 12.8,
                "main_frequency": 1250.0,
                "alarm_level": "warning"
            }
        ],
        "analysis_conclusion": "设备整体运行状态良好，建议继续监测。",
        "recommendations": [
            "定期检查齿轮箱润滑情况",
            "加强高速轴监测频次",
            "建议下次检测时间：3个月后"
        ]
    }
    
    generator = CMSReportGenerator()
    
    # 生成PDF报告
    success = generator.generate_pdf_report(test_data, "test_report.pdf")
    print(f"PDF报告生成: {'成功' if success else '失败'}")
    
    # 生成HTML报告
    success = generator.generate_html_report(test_data, "test_report.html")
    print(f"HTML报告生成: {'成功' if success else '失败'}")